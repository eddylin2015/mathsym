from flask import Flask, render_template, send_file
from docx import Document
from docx.shared import Inches
from io import BytesIO
import io
import qrcode
from docx.shared import Cm
from docx.enum.section import WD_ORIENT #處理文件的直向/橫向
from docx.shared import Pt
from docx.table import _Cell
import datetime 
import copy 

app = Flask(__name__)

def add_qrcode(txt):
    file = qrcode.make('sample text')
    buf = io.BytesIO()
    file.save(buf)
    buf.seek(0)
    #return send_file(buf, mimetype='image/jpeg')
    return buf

def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


@app.route('/')
def index():
    t_=datetime.datetime.now().isoformat()
    return f"""
    <a href="/newdocx?r={t_}">new doc</a>
    <a href="/templdocx?r={t_}">templ doc</a>
    """

@app.route('/templdocx')
def templdocx_router():
    document = Document("secret/lebal_templ_schsu951.docx")
    template = document.tables[0]
    tbl = template._tbl
    for i in range(10):
        # Here we do the copy of the table
        new_tbl = copy.deepcopy(tbl)
        # Then we do the replacement
        # replaceText(document, '<>', 'New value')
        #paragrah = document.add_paragraph()
        paragraph =document.paragraphs[0]
        # After that, we add the previously copied table
        paragraph._p.addnext(new_tbl)    

    for table in document.tables:
        #table.columns[0].width = Cm(1.8)
        #table.columns[1].width = Cm(3)
        hdr_cells = table.rows[0].cells
        file = qrcode.make('sample text')
        buf = io.BytesIO()
        file.save(buf)
        buf.seek(0)
        p=hdr_cells[0].paragraphs[0]
        run=p.add_run()
        run.add_picture(buf, width=Cm(1.8)) # width=Inches(1.25))
        p=hdr_cells[1].paragraphs[0]
        cell_run=p.add_run('FDE-951(210916)\n學生實習使用的手提\n電腦\nMT512')
        cell_run.font.name = '微軟正黑體'
        cell_run.font.size = Pt(8)
        cell_run.bold = True    
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename='report.doc')    




@app.route('/newdocx')
def newdocx_router():
    document = Document()
    section = document.sections[0]
    section.left_margin=Cm(0.1)
    section.right_margin=Cm(0.1)
    section.top_margin=Cm(0.25)
    section.bottom_margin=Cm(0.1)    
    #section.orientation = WD_ORIENT.LANDSCAPE

    new_width, new_height = Cm(5), Cm(2.5)
    section.page_width = new_width
    section.page_height= new_height    
    for i in range(10):
        table = document.add_table(rows=1, cols=2)
        #table.allow_autofit = True
        #table.rows[0].width=Cm(14)
        table.columns[0].width = Cm(1.8)
        table.columns[1].width = Cm(3)

        hdr_cells = table.rows[0].cells
        #set_cell_margins(hdr_cells[0], top=0, start=0, bottom=0, end=0)
        #set_cell_margins(hdr_cells[1], top=0, start=0, bottom=0, end=0)
        file = qrcode.make('sample text')
        buf = io.BytesIO()
        file.save(buf)
        buf.seek(0)
        #p=document.add_picture(buf, width=Cm(1.85)) # width=Inches(1.25))
        #hdr_cells[0].width = Cm(1.8)
        p=hdr_cells[0].paragraphs[0]
        run=p.add_run()
        run.add_picture(buf, width=Cm(1.8)) # width=Inches(1.25))
        #p.alignment=WD_ALIGN_PARAGRAPH.
        #hdr_cells[1].add_paragraph( 'FDE-159(210916)\n學生實習使用的手提\n電腦\nMT512')
        #hdr_cells[1].text= 'FDE-159(210916)\n學生實習使用的手提\n電腦\nMT512'
        #hdr_cells[1].width = Cm(3.2)
        p=hdr_cells[1].paragraphs[0]
        cell_run=p.add_run('FDE-951(210916)\n學生實習使用的手提\n電腦\nMT512')
        cell_run.font.name = '微軟正黑體'
        cell_run.font.size = Pt(8)
        cell_run.bold = True

        #document.add_page_break()
    #document.save('demo.docx')

    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename='report.doc')

@app.route('/qrcode')
def qrcode_router():
    file = qrcode.make('sample text')
    buf = io.BytesIO()
    file.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype='image/jpeg')



from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

if __name__=='__main__':
    app.run(host='0.0.0.0',port=84,debug=True)
    
    #document = Document('schsub159.docx')
    #for block in iter_block_items(document):
    #   print('found one')
    #   print(block.text if isinstance(block, Paragraph) else '<table>')

    #img = qrcode.make('Some data here')
    #type(img)  # qrcode.image.pil.PilImage
    #img.save("some_file.png")    

    """
    
    document.add_heading('Document Title', 0)
    
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')
    
    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    
    document.add_picture('some_file.png', width=Inches(1.25))
    file = qrcode.make('sample text')
    buf = io.BytesIO()
    file.save(buf)
    buf.seek(0)
    document.add_picture(buf, width=Inches(1.25))
    
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )
    
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    
    document.add_page_break()
    
    document.save('demo.docx')
    """
    """
    https://mlhive.com/2022/04/working-with-tables-in-python-docx
    """