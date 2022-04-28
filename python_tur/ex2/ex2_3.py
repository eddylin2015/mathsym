#<div class="small-8 cell infoContentContainer">
#  <h5 class="captionSize">
#     <a href="https://www.gcs.gov.mo/detail/zh-hant/">
#       <span class="infoSubject">教青局向教育界送上節日祝福</span>
#     </a>
#  </h5>
#  <div class="dateNSourceContainer hide-for-small-only grid-x">
#    <div class="sourceLabel shrink cell" title="教育及青年發展局">教育及青年發展局</div>
#    <div class="render_timeago_css" datetime="2021-09-10" data-tid="5">昨天</div>
#  </div>
#  <div class="baseContent baseSize hide-for-small-only fade" 
#      style="max-height:4.5rem;">
#     教育暨青年局與高等教育局於今年初合併為教育及青年發展局後，秉持特區政府“優化教育及人
#     才培養工作”、“深化師資隊伍建設”的施政方針，持續革新和優化教師隊伍建設。9
#  </div>
#</div>
def out_txt(msg):
    print(msg)
    with open(f"./out/news_title.txt", 'a',encoding='utf-8') as the_file:
        the_file.write(str(msg))
        the_file.write('\n')

def out_docx(filename,context):
    from docx import Document
    document = Document()
    document.add_heading(filename, 0)
    document.add_paragraph(
    context, style='List Bullet'
    )
    document.save("./out/"+filename+".docx")

from bs4 import BeautifulSoup
import requests
_url=f"https://www.gcs.gov.mo/list/zh-hant/news/%E6%95%99%E8%82%B2%E9%AB%94%E8%82%B2?7"
title_res = requests.get(_url)
title_res.encoding = "utf-8"
soup = BeautifulSoup(title_res.text, "html.parser")
links = soup.select('[class="small-8 cell infoContentContainer"]')
for link in links:
    news_title=link.select("h5.captionSize")[0].getText().replace("'","").replace('"',"").replace("\n","")
    news_context = link.getText()
    out_txt(news_title)
    out_docx(news_title,news_context)
print("finished .")

