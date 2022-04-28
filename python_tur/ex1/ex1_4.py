from os import listdir
def out_docx(mypath,filename,jpeglist):
    from docx import Document
    document = Document()
    document.add_heading('Document Title', 0)
    for f in jpeglist:
        pic_path=f"{mypath}{f}"
        document.add_picture(pic_path)
    document.save('./out/'+filename)

mypath = "./in/"
files = listdir(mypath)
category={}
for f in files:
    if ".jpg" in f  and '-' in f:
        key , x =f.split('-')
        if key in category:
            category[key].append( f )
        else:
            category[key]=[f]

for k in category:
    filename =f"{k}.docx"
    print(filename)
    print(category[k])
    out_docx(mypath,filename,category[k])

